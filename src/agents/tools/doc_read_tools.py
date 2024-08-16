import os
import fitz  # PyMuPDF
from docx import Document
import openpyxl
from openpyxl.drawing.image import Image
import mimetypes
import hashlib
import chardet
import io
import tempfile

def read_text_file_or_none(file_path):
    """
    读取文件内容，判断是文本文件还是二进制文件。
    
    :param file_path: 文件的路径
    :return: 如果是文本文件，返回文件内容字符串；如果是二进制文件，返回 None。
    """
    def is_text_file(file_path):
        """
        判断文件是否为文本文件。
        
        :param file_path: 文件的路径
        :return: 如果是文本文件返回 True，否则返回 False。
        """
        try:
            with open(file_path, 'rb') as f:
                data = f.read(1024)  # 读取文件的前 1024 字节
            data.decode('utf-8')  # 尝试用 UTF-8 解码
            return True  # 解码成功，返回 True 表示是文本文件
        except UnicodeDecodeError:
            return False  # 解码失败，返回 False 表示是二进制文件

    def detect_encoding(file_path):
        """
        检测文本文件的编码方式。
        
        :param file_path: 文件的路径
        :return: 检测到的文件编码。
        """
        with open(file_path, 'rb') as f:
            data = f.read(10000)  # 读取文件的一部分（前 10000 字节）
        result = chardet.detect(data)  # 使用 chardet 检测编码
        return result['encoding']

    # 判断文件是否为文本文件
    if is_text_file(file_path):
        encoding = detect_encoding(file_path)  # 检测文件编码
        with open(file_path, 'r', encoding=encoding) as f:
            text = f.read()  # 读取整个文件内容
        return text  # 返回文件内容
    else:
        return None  # 返回 None 表示是二进制文件

def get_file_extension(mime_type):
    # 根据MIME类型获取文件扩展名
    return mimetypes.guess_extension(mime_type) or '.bin'
def save_embedded_file(cell, md_files_dir, file_counter):
    for obj in cell._hyperlinks:
        if hasattr(obj, 'target') and hasattr(obj.target, 'content'):
            content = obj.target.content
            mime_type = obj.target.content_type
            extension = get_file_extension(mime_type)
            
            # 使用内容的哈希值作为文件名的一部分，以避免重复
            content_hash = hashlib.md5(content).hexdigest()[:8]
            filename = f"file_{file_counter}_{content_hash}{extension}"
            filepath = os.path.join(md_files_dir, filename)
            
            with open(filepath, 'wb') as f:
                f.write(content)
            
            return filename
    return None

def read_pdf(file_path, tmp_dir):
    """
    读取PDF文件并提取文本内容，同时保存嵌入的图片。

    :param file_path: PDF文件的路径
    :param tmp_dir: 临时目录路径
    :return: 包含PDF文本内容的Markdown字符串和保存的图片文件完整路径列表
    """
    text = ""
    files = []
    pdf_filename = os.path.splitext(os.path.basename(file_path))[0]
    pdf_files_dir = os.path.join(tmp_dir, f"{pdf_filename}_pdf_files")
    os.makedirs(pdf_files_dir, exist_ok=True)

    with fitz.open(file_path) as doc:
        total_pages = len(doc)
        for page_num, page in enumerate(doc):
            text += f"\n\n>[Page {page_num + 1} of {total_pages}]\n\n"
            text += page.get_text()
            text += "\n\n"

            for img_index, img in enumerate(page.get_images(full=True)):
                xref = img[0]
                try:
                    base_image = doc.extract_image(xref)
                except:
                    continue
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                image_filename = f"image_page{page_num + 1}_{img_index + 1}.{image_ext}"
                image_path = os.path.join(pdf_files_dir, image_filename)
                
                with open(image_path, "wb") as img_file:
                    img_file.write(image_bytes)
                
                files.append(image_path)
                text += f"![Image]({os.path.relpath(image_path, tmp_dir)})\n\n"

    return text, files

def read_docx(file_path, tmp_dir):
    """
    读取DOCX文件并提取文本内容，同时保存嵌入的图片和其他文件。

    :param file_path: DOCX文件的路径
    :param tmp_dir: 临时目录路径
    :return: 包含DOCX文本内容的Markdown字符串和保存的文件完整路径列表
    """
    doc = Document(file_path)
    text = ""
    files = []
    docx_filename = os.path.splitext(os.path.basename(file_path))[0]
    docx_files_dir = os.path.join(tmp_dir, f"{docx_filename}_docx_files")
    os.makedirs(docx_files_dir, exist_ok=True)

    for para in doc.paragraphs:
        text += para.text + "\n\n"

    for i, table in enumerate(doc.tables):
        text += f"### 表格 {i+1}\n\n"
        text += "|" + "|".join(cell.text for cell in table.rows[0].cells) + "|\n"
        text += "|" + "|".join("---" for _ in table.rows[0].cells) + "|\n"
        for row in table.rows[1:]:
            text += "|" + "|".join(cell.text for cell in row.cells) + "|\n"
        text += "\n"

    for shape in doc.inline_shapes:
        if shape.type == 3:  # 图片
            image = shape._inline.graphic.graphicData.pic.blipFill.blip
            image_filename = f"image_{len(files) + 1}.png"
            image_path = os.path.join(docx_files_dir, image_filename)
            with open(image_path, "wb") as f:
                f.write(image.embed)
            files.append(image_path)
            text += f"![Image]({os.path.relpath(image_path, tmp_dir)})\n\n"

    return text, files

def save_embedded_file(cell, excel_files_dir, file_counter):
    """
    保存Excel单元格中嵌入的文件。

    :param cell: Excel单元格对象
    :param excel_files_dir: 保存文件的目录
    :param file_counter: 文件计数器
    :return: 保存的文件名，如果没有嵌入文件则返回None
    """
    if cell.hyperlink and cell.hyperlink.target:
        try:
            content = cell.hyperlink.target.content
            mime_type = cell.hyperlink.target.content_type
            extension = mimetypes.guess_extension(mime_type) or '.bin'
            
            content_hash = hashlib.md5(content).hexdigest()[:8]
            filename = f"file_{file_counter}_{content_hash}{extension}"
            filepath = os.path.join(excel_files_dir, filename)
            
            with open(filepath, 'wb') as f:
                f.write(content)
            
            return filename
        except AttributeError:
            # 如果hyperlink没有target属性或target没有content属性，则跳过
            pass
    return None
def read_excel(excel_file, tmp_dir):
    """
    将Excel文件转换为Markdown格式，并保存嵌入的图片和文件。

    :param excel_file: Excel文件的路径
    :param tmp_dir: 临时目录路径
    :return: 包含Excel内容的Markdown字符串和保存的文件完整路径列表
    """
    workbook = openpyxl.load_workbook(excel_file, data_only=True)
    markdown = ""
    files = []
    file_counter = 1

    excel_filename = os.path.splitext(os.path.basename(excel_file))[0]
    excel_files_dir = os.path.join(tmp_dir, f"{excel_filename}_excel_files")
    os.makedirs(excel_files_dir, exist_ok=True)

    def process_cell_value(value):
        if isinstance(value, str):
            # 替换换行符为 \n，并替换竖线为转义的竖线
            return value.replace('\n', '\\n').replace('|', '\\|')
        elif value is not None:
            return str(value).replace('|', '\\|')
        return ""

    for sheet in workbook.sheetnames:
        worksheet = workbook[sheet]
        markdown += f"# {sheet}\n\n"

        merged_cells = worksheet.merged_cells.ranges

        # 处理该工作表中的所有图片
        image_counter = 1
        for image in worksheet._images:
            img_format = image.format if hasattr(image, 'format') else 'png'
            img_filename = f"image_{sheet}_{image_counter}.{img_format}"
            img_path = os.path.join(excel_files_dir, img_filename)
            
            # 从BytesIO对象中读取字节数据
            image_data = image.ref.getvalue() if isinstance(image.ref, io.BytesIO) else image.ref
            
            with open(img_path, "wb") as img_file:
                img_file.write(image_data)
            
            files.append(img_path)
            image_counter += 1

        for row in worksheet.iter_rows():
            row_md = "|"
            for cell in row:
                is_merged = any(cell.coordinate in mr for mr in merged_cells)
                
                if is_merged:
                    if cell == worksheet.cell(cell.row, cell.column):
                        value = process_cell_value(cell.value)
                        row_md += f" {value} |"
                    else:
                        row_md += " |"
                else:
                    value = process_cell_value(cell.value)
                    row_md += f" {value} |"

                # 检查该单元格是否有图片
                for image in worksheet._images:
                    if image.anchor._from.col == cell.column - 1 and image.anchor._from.row == cell.row - 1:
                        img_format = image.format if hasattr(image, 'format') else 'png'
                        img_filename = f"image_{sheet}_{worksheet._images.index(image) + 1}.{img_format}"
                        img_path = os.path.join(excel_files_dir, img_filename)
                        row_md += f" ![Image]({os.path.relpath(img_path, tmp_dir)}) |"

                embedded_file = save_embedded_file(cell, excel_files_dir, file_counter)
                if embedded_file:
                    file_path = os.path.join(excel_files_dir, embedded_file)
                    files.append(file_path)
                    row_md += f" [Embedded File]({os.path.relpath(file_path, tmp_dir)}) |"
                    file_counter += 1

            markdown += row_md + "\n"

            if row[0].row == 1:
                markdown += "|" + " --- |" * len(row) + "\n"

        markdown += "\n\n"

    return markdown, files

def file2text(file_path, tmp_dir=None):
    """
    将各种文件格式转换为文本或Markdown格式。

    :param file_path: 要读取的文件路径
    :param tmp_dir: 临时目录路径，如果为None则使用系统临时目录
    :return: 包含文件内容的文本或Markdown字符串、保存的文件完整路径列表（如果有的话）和转换后的Markdown文件路径
    """
    if tmp_dir is None:
        tmp_dir = tempfile.gettempdir()

    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()

    conversion_notice = f"""
> 本文档是由{file_extension[1:].upper()}文件转换为Markdown格式生成的。
> 在转换过程中可能会出现断行错误和混乱，排版、标题、字体、颜色等格式的变化，以及内容丢失，从而导致文档的可读性、准确性和完整性受损。  
> 阅读中你需要克服这些问题。
---

"""

    # 保留原始文件的目录结构
    relative_path = os.path.relpath(file_path, start=os.path.dirname(file_path))
    output_md_file_path = os.path.join(tmp_dir, relative_path)
    output_md_file_path = os.path.splitext(output_md_file_path)[0] + "_converted.md"

    # 确保目标目录存在
    os.makedirs(os.path.dirname(output_md_file_path), exist_ok=True)

    if file_extension == '.pdf':
        content, files = read_pdf(file_path, tmp_dir)
    elif file_extension in ['.docx', '.doc']:
        content, files = read_docx(file_path, tmp_dir)
    elif file_extension in ['.xlsx', '.xls']:
        content, files = read_excel(file_path, tmp_dir)
    else:
        # 尝试用 read_text_file_or_none 读取
        content = read_text_file_or_none(file_path)
        if content is not None:
            return content, [], file_path  # 返回文本内容和空的文件列表
        else:
            return None, [], file_path

    with open(output_md_file_path, 'w', encoding='utf-8') as f:
        f.write(conversion_notice + content)

    return conversion_notice + content, files, output_md_file_path


if __name__ == '__main__':
    # 测试 file2text 函数
    test_files = ['test0.docx', 'test1.pdf', 'test2.pdf', 'test3.xlsx']
    
    for file in test_files:
        try:
            content, files = file2text(file)
            output_file = f"{os.path.splitext(file)[0]}_output.md"
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(content)
            print(f"Successfully processed {file}. Output saved to {output_file}")
            if files:
                print(f"Associated files: {files}")
        except Exception as e:
            print(f"Error processing {file}: {str(e)}")